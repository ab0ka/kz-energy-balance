#!/usr/bin/env python3
"""
Build the methodology document explaining how the article was written.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(__file__)
OUT = os.path.join(BASE, 'Methodology_How_Article_Was_Written.docx')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(12)


def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def subheading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(12)


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(12)


# ── Document content ──

title('Methodology Document:\nHow the Scientific Article Was Written')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sekenov Gabit | March 2026')
r.font.size = Pt(11)
r.italic = True
p.paragraph_format.space_after = Pt(16)

# ── 1 ──
heading('1. Overview')

body(
    'This document describes the complete methodology used to research, write, and produce the '
    'scientific article titled "AI Data Center Infrastructure in Kazakhstan: Technical Resource '
    'Analysis and Site Optimization" for submission to the International Journal of Digitalization '
    '(IJDIGIT). The document covers the research process, data sources, analytical methods, '
    'figure generation, and quality assurance steps.'
)

# ── 2 ──
heading('2. Research Process')

subheading('2.1 Problem Identification')
body(
    'The original draft of the article was identified as requiring a complete rewrite due to '
    'insufficient depth of analysis, limited referencing, and lack of professional academic formatting. '
    'The core research question was preserved: evaluating the technical and economic feasibility of '
    'deploying AI data center infrastructure in Kazakhstan.'
)

subheading('2.2 Literature Review')
body('The following categories of sources were systematically reviewed:')
bullet('Academic papers on data center energy efficiency (IEEE, Elsevier, Nature portfolio journals)')
bullet('Zhakiyev et al. (2023) — TIMES model for Kazakhstan decarbonization to 2060 (Energies, vol. 16)')
bullet('Zhakiyev et al. (2024) — PyPSA-KZ power system model for coal exit scenarios (Engineered Science, vol. 29)')
bullet('IEA reports on global data center energy consumption and AI energy demand')
bullet('NVIDIA technical documentation on GPU power density (H100, GB200, Rubin platforms)')
bullet('Government policy documents: 2026 "Year of Digitalization and AI" decree, tariff freeze order')
bullet('World Bank Water Security Assessment for Kazakhstan (2024)')
bullet('ASHRAE TC 9.9 Thermal Guidelines for Data Processing Environments')
bullet('Uptime Institute Global Data Center Survey (2024)')

subheading('2.3 Data Collection')
body('Quantitative data was collected from the following sources:')
bullet('Meteorological data for Astana (2015-2025): monthly temperature averages, extremes, and humidity')
bullet('Electricity tariff data: Ministry of Energy of Kazakhstan, GlobalPetrolPrices.com, industry reports')
bullet('Water resource data: FAO AQUASTAT, Earth.Org, World Bank assessments')
bullet('Data center PUE benchmarks: Google, AWS, Microsoft sustainability reports, Uptime Institute surveys')
bullet('Renewable energy capacity: QazaqGreen, GlobalData, The Astana Times reporting')
bullet('Grid carbon intensity: Ember Climate, Low Carbon Power, Climate Analytics databases')

# ── 3 ──
heading('3. Analytical Methods')

subheading('3.1 Thermal Modeling (PUE)')
body(
    'The PUE model uses a piecewise function of ambient temperature with three regimes: '
    '(1) Full free cooling (T < 10 deg C), PUE contribution from cooling = 0.04-0.06; '
    '(2) Partial economization (10-20 deg C), linear interpolation of mechanical cooling load; '
    '(3) Full mechanical cooling (T > 20 deg C), PUE contribution = 0.20-0.26. '
    'Monthly PUE values were calculated using Astana mean temperatures, then annualized '
    'as a weighted average. The 10 deg C threshold follows ASHRAE TC 9.9 Class A2 guidelines '
    'for direct airside economization.'
)

subheading('3.2 Economic Analysis')
body(
    'OPEX was calculated for a reference 35 MW facility (30.6 MW IT load at PUE 1.15) '
    'operating continuously (8,760 hours/year). Tariff data from six global regions was '
    'used for comparison. The annual cost formula: E_annual = P_total x 8,760 hours x tariff. '
    'Two Kazakhstan scenarios were modeled: grid-connected ($0.076/kWh) and co-location '
    'at generation ($0.031/kWh). The tariff freeze through 2032 was treated as a fixed '
    'assumption based on the December 2025 Ministry of Energy order.'
)

subheading('3.3 Water Consumption Analysis')
body(
    'Water consumption was modeled for five cooling technologies using published data from '
    'Mytton (2021) in npj Clean Water and supplemented by manufacturer specifications. '
    'The reference metric is annual water consumption in cubic meters for a 35 MW facility. '
    'Sustainability thresholds were derived from World Bank projections of Kazakhstan water '
    'availability trajectories.'
)

subheading('3.4 Multi-Criteria Site Selection')
body(
    'A multi-criteria decision analysis (MCDA) framework was developed with five criteria: '
    'PUE Efficiency (25%), Power Cost (25%), Water Access (20%), Network Connectivity (15%), '
    'and Renewable Energy Potential (15%). Three candidate regions (Astana, Pavlodar/Ekibastuz, '
    'Almaty) were scored on a 1-10 scale. Quantitative scores were derived from the thermal '
    'and economic models; connectivity and renewable potential scores incorporated qualitative '
    'assessment based on infrastructure availability and geographic factors.'
)

# ── 4 ──
heading('4. Figure Generation')

body('All 8 figures were generated programmatically using Python 3 with the matplotlib library. '
     'The generation script (generate_figures.py) produces publication-quality PNG files at 300 DPI.')

bullet('Fig. 1: Global DC Energy Consumption Forecast — grouped bar chart (IEA data, 2020-2030)')
bullet('Fig. 2: Rack Power Density Evolution — bar chart (NVIDIA specs, industry projections)')
bullet('Fig. 3: Astana Climate Profile — line chart with shaded free cooling zone (meteorological data)')
bullet('Fig. 4: Monthly PUE Estimation — line chart comparing Astana vs Nordic vs US average (modeled)')
bullet('Fig. 5: Electricity OPEX Comparison — bar chart across 6 global regions (tariff data)')
bullet('Fig. 6: Water Consumption by Cooling Technology — horizontal log-scale bar chart (Mytton 2021)')
bullet('Fig. 7: Kazakhstan Energy Mix — dual pie chart, 2024 vs 2035 (Ember + Zhakiyev PyPSA-KZ)')
bullet('Fig. 8: Regional Site Selection — radar/spider chart for 3 regions x 5 criteria (MCDA model)')

body(
    'Color scheme uses a professional academic palette (deep blue primary, red accent, green for positive '
    'indicators). All figures include axis labels, legends, annotations, and source attributions.'
)

# ── 5 ──
heading('5. Paper Structure and Formatting')

body('The paper follows the IJDIGIT author guidelines:')
bullet('Page format: A4, margins 15mm sides / 24mm top-bottom')
bullet('Font: Times New Roman, 10pt body text, 16pt title')
bullet('Reference style: IEEE numbered format with square brackets')
bullet('Total length: approximately 10 pages with figures and tables (within 12-page limit)')
bullet('Sections: Abstract, Keywords, Introduction, Literature Review, Methodology, Results and Analysis, Discussion, Strategic Recommendations, Conclusion, Author Contributions, Acknowledgements, Conflict of Interest, References')

body(
    'The paper was assembled programmatically using python-docx to ensure consistent formatting '
    'throughout. The build script (build_paper.py) creates the final DOCX file with embedded figures '
    'and formatted tables.'
)

# ── 6 ──
heading('6. Key Improvements Over Original Draft')

bullet('Added comprehensive Literature Review section with context from Zhakiyev et al. (2023, 2024)')
bullet('Expanded methodology from brief description to rigorous multi-component framework')
bullet('Added 8 professionally generated data visualizations (original had placeholder figures)')
bullet('Increased reference count from 11 to 18 with verified, citable sources')
bullet('Added quantitative analysis: monthly PUE model, OPEX comparison across 6 regions')
bullet('Added water resource constraint analysis with technology comparison table')
bullet('Added multi-criteria site selection framework with radar chart visualization')
bullet('Added Discussion section with critical evaluation of limitations and risks')
bullet('Added Strategic Recommendations with evidence-based policy proposals')
bullet('Added Author Contributions, Acknowledgements, and Conflict of Interest sections per IJDIGIT requirements')
bullet('Improved academic language throughout, replacing informal/vague statements with precise, referenced claims')

# ── 7 ──
heading('7. Source Files')

body('The following files were produced during this work:')
bullet('generate_figures.py — Python script generating all 8 figures as PNG files')
bullet('build_paper.py — Python script assembling the DOCX paper with embedded figures')
bullet('figures/ — directory containing all 8 figure PNG files at 300 DPI')
bullet('Sekenov_AI_DataCenter_Kazakhstan_IJDIGIT.docx — the final paper')
bullet('Methodology_How_Article_Was_Written.docx — this document')

# ── 8 ──
heading('8. References Used in This Process')

refs = [
    'Zhakiyev, N. et al. (2023). Optimization modelling of the decarbonization scenario of the total energy system of Kazakhstan until 2060. Energies, 16(13), 5142.',
    'Zhakiyev, N. et al. (2024). Comprehensive scenario analyses for coal exit and renewable energy development planning of Kazakhstan using PyPSA-KZ. Engineered Science, 29, 1085.',
    'International Energy Agency (2025). Energy and AI. IEA, Paris.',
    'Mytton, D. (2021). Data centre water consumption. npj Clean Water, 4(11).',
    'Ahmed, K.M.U., Bollen, M.H., and Alvarez, M. (2021). A review of data centers energy consumption and reliability modeling. IEEE Access, 9.',
    'Koot, M. and Wijnhoven, F. (2021). Usage impact on data center electricity needs. Applied Energy, 291.',
    'Uptime Institute (2024). Global Data Center Survey Results 2024.',
    'ASHRAE TC 9.9 (2021). Thermal Guidelines for Data Processing Environments, 5th ed.',
    'World Bank (2024). Kazakhstan Water Security Assessment.',
    'NVIDIA Corporation (2026). Rubin Platform Technical Specifications.',
    'Google (2024). Data Centers — Efficiency. Available: https://datacenters.google/efficiency/',
    'Ember Climate (2025). Kazakhstan Electricity Data. Available: https://ember-energy.org/countries-and-regions/kazakhstan/',
    'Low Carbon Power (2025). Kazakhstan Generation and Carbon Intensity.',
    'QazaqGreen (2025). Maximum Electricity Tariffs 2026-2032.',
    'The Astana Times (2026). Ekibastuz Data Center Valley; Year of Digitalization reporting.',
    'Earth.Org (2025). Kazakhstan Water Crisis, Explained with Data.',
    'GlobalData (2025). Kazakhstan Renewable Power Capacity to Reach 12.9 GW by 2035.',
    'Orda.kz (2025). Kazakhstan Proposes Seven-Year Freeze on Electricity Rate Caps.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. {ref}')
    r.font.size = Pt(11)

doc.save(OUT)
print(f'Methodology document saved to: {OUT}')
