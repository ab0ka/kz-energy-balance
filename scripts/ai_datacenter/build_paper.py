#!/usr/bin/env python3
"""
Build the IJDIGIT paper as a DOCX:
"AI Data Center Infrastructure in Kazakhstan: Technical Resource Analysis and Site Optimization"
Author: Sekenov Gabit
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(__file__)
FIGS = os.path.join(BASE, 'figures')
OUT = os.path.join(BASE, 'Sekenov_AI_DataCenter_Kazakhstan_IJDIGIT.docx')

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)


def add_author_info():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Sekenov Gabit')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        'Department of Computational and Data Science',
        'Astana IT University',
        'Astana, Republic of Kazakhstan',
    ]
    for line in lines:
        run = p2.add_run(line + '\n')
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.italic = True
    p2.paragraph_format.space_after = Pt(12)


def add_heading_styled(text, level=1):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.63)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p


def add_body_no_indent(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p


def add_figure(filename, caption, width=Inches(5.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(FIGS, filename), width=width)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = cap.add_run(caption.split('.')[0] + '. ')
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.name = 'Times New Roman'
    rest = '.'.join(caption.split('.')[1:]).strip()
    if rest:
        run_text = cap.add_run(rest)
        run_text.font.size = Pt(9)
        run_text.font.name = 'Times New Roman'
    cap.paragraph_format.space_after = Pt(8)


def add_table(headers, rows, caption):
    # Table caption above
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = cap.add_run(caption.split('.')[0] + '. ')
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.name = 'Times New Roman'
    rest = '.'.join(caption.split('.')[1:]).strip()
    if rest:
        run_t = cap.add_run(rest)
        run_t.font.size = Pt(9)
        run_t.font.name = 'Times New Roman'
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer


# ════════════════════════════════════════════════════════════
# BUILD PAPER
# ════════════════════════════════════════════════════════════

# ── TITLE ──
add_title('AI Data Center Infrastructure in Kazakhstan:\nTechnical Resource Analysis and Site Optimization')
add_author_info()

# ── ABSTRACT ──
p_abs = doc.add_paragraph()
run_abs_label = p_abs.add_run('Abstract')
run_abs_label.bold = True
run_abs_label.font.size = Pt(10)
run_abs_label.font.name = 'Times New Roman'
run_abs_dash = p_abs.add_run('—')
run_abs_dash.bold = True
run_abs_dash.font.size = Pt(10)

abstract_text = (
    "The rapid proliferation of artificial intelligence (AI) has created unprecedented demands on data center "
    "infrastructure, with modern GPU clusters requiring 40–240 kW per rack—an order of magnitude increase over "
    "traditional computing facilities. This paper presents a comprehensive techno-economic analysis of deploying "
    "AI-optimized data center infrastructure in the Republic of Kazakhstan, a nation positioned at the nexus of "
    "abundant energy resources, favorable climatic conditions, and ambitious digitalization targets under the "
    "2026 \"Year of Digitalization and AI\" initiative. We develop a multi-factor assessment framework incorporating "
    "thermal modeling based on decade-long meteorological records (2015–2025), total cost of ownership (TCO) analysis "
    "leveraging the recently enacted 2026–2032 tariff stabilization policy, and water resource constraint modeling "
    "under projected scarcity scenarios. Our thermal analysis demonstrates that Astana's continental climate enables "
    "6–7 months of full free cooling (ambient temperature <10°C), achieving an annualized Power Usage Effectiveness "
    "(PUE) of 1.12—competitive with leading Nordic facilities. Economic modeling reveals that co-location at "
    "generation stations in northern Kazakhstan reduces electricity operational expenditures (OPEX) by 59–63% "
    "compared to U.S. industrial averages, yielding annual savings of $14–16 million for a reference 35 MW facility. "
    "However, we identify water scarcity as the dominant constraint: traditional evaporative cooling would consume "
    "approximately 900,000 m³ annually per 35 MW facility, which is unsustainable given Kazakhstan's projected 50% "
    "water shortfall by 2040. We propose a regional site selection matrix evaluating three candidate regions across "
    "five weighted criteria and recommend closed-loop liquid cooling as a mandatory design requirement for sustainable "
    "deployment."
)
run_abs_text = p_abs.add_run(abstract_text)
run_abs_text.font.size = Pt(10)
run_abs_text.font.name = 'Times New Roman'
run_abs_text.italic = True
p_abs.paragraph_format.space_after = Pt(6)

# ── KEYWORDS ──
p_kw = doc.add_paragraph()
run_kw_label = p_kw.add_run('Keywords')
run_kw_label.bold = True
run_kw_label.font.size = Pt(10)
run_kw_label.font.name = 'Times New Roman'
run_kw = p_kw.add_run('—AI Data Centers, Power Usage Effectiveness, Free Cooling, Water Scarcity, Renewable Energy, Kazakhstan.')
run_kw.font.size = Pt(10)
run_kw.font.name = 'Times New Roman'
run_kw.italic = True
p_kw.paragraph_format.space_after = Pt(10)

# ════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ════════════════════════════════════════════════════════════
add_heading_styled('I. Introduction')

add_body(
    "The global proliferation of artificial intelligence has fundamentally reshaped the computational "
    "infrastructure landscape. Unlike traditional enterprise workloads characterized by cyclical utilization "
    "patterns, AI training clusters operate at sustained maximum capacity for weeks to months, generating "
    "unprecedented power density and thermal loads [1]. The training of a single large language model (LLM) "
    "can consume on the order of 50–100 GWh of electricity [2], while inference workloads, though individually "
    "lighter, contribute cumulatively to an estimated global data center electricity consumption of 415 TWh in "
    "2024—approximately 1.5% of global electricity generation—with projections reaching 650–1,050 TWh by 2026 "
    "and 945 TWh by 2030 [3]."
)

add_body(
    "The hardware driving this transformation has undergone a parallel revolution in power density. "
    "Traditional CPU-based server racks operated at 5–15 kW, well within the design parameters of conventional "
    "air-cooled facilities. Modern AI accelerators—exemplified by the NVIDIA H100 (700 W) and the GB200 NVL72 "
    "system (132 kW per rack)—have fundamentally exceeded these limits, with projected rack densities approaching "
    "500 kW by 2027 [4]. This shift necessitates a paradigm change in power distribution, cooling systems, and "
    "facility design."
)

add_body(
    "The Republic of Kazakhstan has emerged as a strategically significant candidate for AI infrastructure "
    "deployment. In early 2026, President Tokayev declared the year as the \"Year of Digitalization and AI,\" "
    "accompanied by a comprehensive action plan encompassing 8 priority blocks and 10 key initiatives [5]. "
    "The government has initiated the Ekibastuz Data Center Valley project, envisioning a 1 GW computing "
    "hub with up to $30 billion in targeted investment [6]. Kazakhstan possesses several structural advantages: "
    "an extreme continental climate enabling extensive free cooling, electricity tariffs frozen through 2032 at "
    "rates substantially below Western benchmarks [7], and a renewable energy potential exceeding 760 GW for "
    "wind alone [8]."
)

add_body(
    "However, realization of this potential faces significant challenges. Kazakhstan's electricity grid "
    "remains predominantly coal-fired (54% of generation in 2024), with a carbon intensity of approximately "
    "600 gCO₂eq/kWh—35% above the global average [9]. Infrastructure aging is severe, with 50–70% physical "
    "wear across the thermal power plant fleet [10]. Most critically, the nation confronts an acute water "
    "deficit: only 42% of theoretical water resources are technically accessible, and projections indicate a "
    "50% shortfall by 2040 [11]. These constraints demand a systematic, evidence-based approach to site "
    "selection and facility design."
)

add_body(
    "This paper contributes a rigorous techno-economic framework for evaluating AI data center deployment "
    "in Kazakhstan, integrating thermal modeling, economic analysis, and environmental constraint assessment. "
    "We draw upon recent energy system modeling by Zhakiyev et al. [10], [12] to contextualize our findings "
    "within Kazakhstan's broader energy transition trajectory."
)

# ════════════════════════════════════════════════════════════
# II. LITERATURE REVIEW
# ════════════════════════════════════════════════════════════
add_heading_styled('II. Literature Review')

add_body(
    "The academic literature on data center energy efficiency has expanded rapidly in response to the "
    "AI-driven surge in power demand. Ahmed et al. [1] provide a comprehensive review of data center energy "
    "consumption models, establishing the baseline understanding of facility-level power distribution. "
    "Koot and Wijnhoven [2] analyze the relationship between workload characteristics and energy consumption, "
    "demonstrating that AI workloads exhibit fundamentally different energy profiles compared to traditional "
    "computing tasks. Mytton [13] quantifies data center water consumption, establishing the framework of "
    "Water Usage Effectiveness (WUE) that informs our analysis."
)

add_body(
    "In the context of Kazakhstan's energy system, Zhakiyev et al. [12] developed an optimization model "
    "using the TIMES framework to evaluate decarbonization pathways for Kazakhstan's total energy system through "
    "2060. Their carbon neutrality scenario projects electricity generation growing to 472 TWh by 2060 (nearly "
    "5× the 2017 baseline), with renewable energy achieving a 50% share. The study identifies an additional "
    "$544 billion in required investment beyond the reference scenario, highlighting the scale of the energy "
    "transition challenge. In a subsequent study, Zhakiyev et al. [10] developed PyPSA-KZ, a spatially resolved "
    "power system model based on the Python for Power System Analysis framework. This model provides granular "
    "data on Kazakhstan's grid topology, revealing critical spatial imbalances: generation capacity is "
    "concentrated in the north, while the south faces a structural deficit, and the western zone remains "
    "physically disconnected from the main grid. Their analysis of a 30% RES scenario for 2035 estimates "
    "annual system costs of EUR 5.08 billion and identifies optimal regions for wind (central and eastern "
    "Kazakhstan) and solar (Turkestan, Almaty, Aktobe) deployment."
)

add_body(
    "The intersection of data center siting and national energy policy remains underexplored in the "
    "Central Asian context. While Northern European countries have been extensively studied as free cooling "
    "destinations [14], Kazakhstan's unique combination of extreme continental climate, rapidly evolving "
    "energy policy, and acute water constraints has not been systematically evaluated for high-density AI "
    "infrastructure. This paper addresses that gap."
)

# ════════════════════════════════════════════════════════════
# III. METHODOLOGY
# ════════════════════════════════════════════════════════════
add_heading_styled('III. Methodology')

add_body(
    "Our assessment framework integrates three analytical components: thermal modeling, economic "
    "analysis, and environmental constraint evaluation. Each component utilizes publicly available data "
    "sources and established engineering methodologies."
)

add_heading_styled('A. Thermal Modeling', level=2)

add_body(
    "We model the Power Usage Effectiveness (PUE) as a function of ambient temperature using historical "
    "meteorological data for Astana, Kazakhstan (2015–2025). PUE is defined as the ratio of total facility "
    "power to IT equipment power:"
)

p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_eq = p_eq.add_run('PUE = (P_IT + P_Cooling + P_Electrical) / P_IT       (1)')
run_eq.font.size = Pt(10)
run_eq.font.name = 'Times New Roman'
run_eq.italic = True
p_eq.paragraph_format.space_before = Pt(6)
p_eq.paragraph_format.space_after = Pt(6)

add_body(
    "where P_Cooling is modeled as a piecewise function of ambient temperature (T_amb). For T_amb < 10°C, "
    "we assume full free cooling with compressor load approximately zero, yielding a PUE contribution of "
    "0.04–0.06. For 10°C ≤ T_amb ≤ 20°C, partial economization is applied with a linear interpolation "
    "of mechanical cooling load. For T_amb > 20°C, full mechanical cooling is required. The 10°C threshold "
    "is consistent with ASHRAE TC 9.9 Class A2 guidelines for direct airside economization [15]."
)

add_heading_styled('B. Economic Model', level=2)

add_body(
    "Operational expenditure (OPEX) for electricity is calculated using validated industrial tariff data "
    "from the Ministry of Energy of Kazakhstan. We evaluate two deployment scenarios: (1) Grid-Connected, "
    "utilizing standard industrial tariffs at 36.66 KZT/kWh ($0.076 USD), applicable for facilities in "
    "Astana or Almaty; and (2) Co-Location at generation, with direct access to station-bus pricing of "
    "12–15 KZT/kWh ($0.025–$0.031 USD), applicable for facilities adjacent to Ekibastuz GRES-1 or GRES-2 "
    "power stations [7]. The tariff stability guaranteed through the 2026–2032 freeze order provides a "
    "unique advantage for long-term TCO modeling, eliminating the tariff volatility risk inherent in most "
    "markets."
)

add_heading_styled('C. Water Resource Assessment', level=2)

add_body(
    "Water consumption is modeled for five cooling technologies based on published data from Mytton [13] "
    "and manufacturer specifications. We evaluate annual water consumption for a reference 35 MW facility "
    "and assess sustainability against Kazakhstan's projected water availability trajectories, utilizing "
    "data from the World Bank Water Security Assessment [16] and the Ministry of Ecology's 2024 Water Code "
    "provisions."
)

add_heading_styled('D. Regional Site Selection', level=2)

add_body(
    "We develop a multi-criteria decision analysis (MCDA) framework evaluating three candidate regions—"
    "Astana, Pavlodar/Ekibastuz, and Almaty—across five weighted criteria: PUE Efficiency (25%), Power "
    "Cost (25%), Water Access (20%), Network Connectivity (15%), and Renewable Energy Potential (15%). "
    "Scores are assigned on a 1–10 scale based on quantitative data where available, supplemented by "
    "qualitative assessment for connectivity metrics."
)

# ════════════════════════════════════════════════════════════
# IV. RESULTS AND ANALYSIS
# ════════════════════════════════════════════════════════════
add_heading_styled('IV. Results and Analysis')

# ── A. Global Context ──
add_heading_styled('A. Global Data Center Energy Demand Context', level=2)

add_body(
    "To contextualize Kazakhstan's opportunity, we first examine the global trajectory of data center "
    "energy demand. As illustrated in Fig. 1, the International Energy Agency projects global data center "
    "electricity consumption to grow from 415 TWh in 2024 to approximately 945 TWh by 2030, representing "
    "a compound annual growth rate (CAGR) of approximately 15% [3]. AI workloads, currently accounting for "
    "an estimated 60 TWh (14% of total), are projected to grow disproportionately, reaching approximately "
    "472 TWh (50%) by 2030. This growth is driving a global search for regions offering favorable economics "
    "and climatic conditions."
)

add_figure('fig1_global_dc_energy.png',
           'Fig. 1. Global data center energy consumption forecast (2020–2030), showing the accelerating share '
           'of AI-specific workloads. Data synthesized from IEA [3] and industry projections.')

# ── B. Power Density Requirements ──
add_heading_styled('B. Power Density Requirements for AI Infrastructure', level=2)

add_body(
    "The evolution of rack power density, depicted in Fig. 2, represents a fundamental challenge for "
    "facility design. Traditional data centers were engineered for 5–15 kW per rack, with power distribution "
    "and cooling systems dimensioned accordingly. The introduction of GPU-accelerated computing has driven "
    "a dramatic escalation: NVIDIA's H100-based systems operate at approximately 70 kW per rack, while the "
    "GB200 NVL72 platform reaches 132 kW [4]. Industry roadmaps project rack densities of 240–500 kW by "
    "2027, with 1 MW per rack on longer-term product roadmaps."
)

add_figure('fig2_rack_density.png',
           'Fig. 2. Evolution of data center rack power density from traditional CPU-based systems to '
           'projected AI GPU clusters, illustrating the order-of-magnitude increase requiring fundamental '
           'redesign of power and cooling infrastructure.')

add_body(
    "For a reference facility housing 30,000 GPUs (NVIDIA H100 class), the critical IT power requirement "
    "is approximately 30.6 MW. Assuming an achievable PUE of 1.15 in Kazakhstan's climate, the total "
    "facility load reaches 35.2 MW, translating to an annual energy consumption of:"
)

p_eq2 = doc.add_paragraph()
p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_eq2 = p_eq2.add_run('E_annual = 35.2 MW × 8,760 h ≈ 308,352 MWh       (2)')
run_eq2.font.size = Pt(10)
run_eq2.font.name = 'Times New Roman'
run_eq2.italic = True
p_eq2.paragraph_format.space_before = Pt(6)
p_eq2.paragraph_format.space_after = Pt(6)

add_body(
    "This consumption level is comparable to the electricity demand of a city of approximately 150,000 "
    "residents, underscoring the infrastructure-scale implications of AI facility deployment."
)

# ── C. Climate Analysis and PUE Optimization ──
add_heading_styled('C. Climate Analysis and PUE Optimization', level=2)

add_body(
    "Astana's extreme continental climate presents a significant advantage for data center cooling. "
    "Fig. 3 illustrates the annual temperature profile, with mean temperatures ranging from −15°C in "
    "January to +20°C in July. Critically, ambient temperatures remain below the 10°C free cooling "
    "threshold for approximately 7 months (October through April), enabling full bypass of mechanical "
    "cooling systems during this period. Even during the warmest months (June–August), peak temperatures "
    "rarely exceed the ASHRAE A1 upper limit of 27°C for extended periods."
)

add_figure('fig3_climate_profile.png',
           'Fig. 3. Astana annual climate profile showing monthly temperature ranges and the '
           'free cooling window. Ambient temperatures below 10°C (shaded zone) enable compressor-free '
           'cooling for approximately 7 months annually.')

add_body(
    "Fig. 4 presents our month-by-month PUE model for a facility in Astana, benchmarked against "
    "Nordic data centers and the U.S. industry average. During the winter months (November–March), "
    "the modeled PUE reaches 1.04–1.05, driven entirely by fan power for airside economization. "
    "The peak PUE of 1.26 occurs in July, when mechanical cooling is required. The annualized PUE "
    "of 1.12 is highly competitive with Nordic benchmarks (1.07–1.16) and substantially below the "
    "U.S. industry average of 1.45 [17]. This result is consistent with Google's best-in-class fleet "
    "PUE of 1.09, suggesting that Astana-based facilities can approach hyperscaler efficiency levels."
)

add_figure('fig4_pue_monthly.png',
           'Fig. 4. Monthly PUE estimation for Astana compared to Nordic benchmarks and the U.S. industry '
           'average. The Astana facility achieves an annualized PUE of 1.12, driven by extensive free cooling '
           'during October–April.')

# ── D. Economic Analysis ──
add_heading_styled('D. Economic Analysis (OPEX)', level=2)

add_body(
    "A critical competitive advantage for Kazakhstan is the combination of inherently low electricity "
    "prices and the regulatory certainty provided by the 2026–2032 tariff freeze. In December 2025, the "
    "Ministry of Energy approved maximum tariff caps frozen at 2025 levels through 2032, following a "
    "presidential directive on inflation management [7]. This seven-year price guarantee is unprecedented "
    "among potential data center markets and enables high-confidence long-term financial modeling."
)

# Table I: OPEX Comparison
add_table(
    headers=['Region', 'Tariff ($/kWh)', 'Annual OPEX (M$)', 'Savings vs. US'],
    rows=[
        ['Germany', '$0.180', '$55.2', '−116%'],
        ['Singapore', '$0.120', '$36.8', '−44%'],
        ['US (Industrial Avg.)', '$0.083', '$25.6', '—'],
        ['Kazakhstan (Grid)', '$0.076', '$23.4', '8.6%'],
        ['Iceland', '$0.045', '$13.8', '46%'],
        ['Kazakhstan (Direct)', '$0.031', '$9.5', '62.9%'],
    ],
    caption='Table I. Annual Electricity OPEX Comparison for a 35 MW AI Data Center Facility'
)

add_body(
    "As shown in Table I and Fig. 5, co-location at generation stations in the Ekibastuz region "
    "offers the most favorable economics globally, with annual electricity OPEX of $9.5 million—63% "
    "below the U.S. industrial average and 83% below German rates. Even grid-connected facilities in "
    "Kazakhstan achieve modest savings of 8.6% versus the U.S. baseline. Notably, Kazakhstan's direct "
    "generation pricing ($0.031/kWh) undercuts even Iceland ($0.045/kWh), which has been the benchmark "
    "for low-cost data center power."
)

add_figure('fig5_opex_comparison.png',
           'Fig. 5. Annual electricity OPEX comparison for a 35 MW facility across global markets. '
           'Co-location in Kazakhstan at generation stations provides the lowest cost globally, with '
           '62.9% savings versus the U.S. industrial average.')

# ── E. Water Resource Constraints ──
add_heading_styled('E. Water Resource Constraints', level=2)

add_body(
    "Water scarcity represents the most critical environmental constraint for data center deployment "
    "in Kazakhstan. The nation's total renewable water resources are estimated at 108 km³/year, of which "
    "only 42% (approximately 45 km³) is technically accessible due to infrastructure limitations [11]. "
    "Per capita water availability has declined 21% since 1999, and projections indicate a potential 50% "
    "shortfall relative to national demand by 2040 [16]. The situation is particularly acute in the western "
    "and southern regions, where 70% of the territory may face severe water stress by 2050."
)

add_body(
    "Fig. 6 presents a logarithmic comparison of annual water consumption across five cooling "
    "technologies for the reference 35 MW facility. Traditional evaporative cooling towers would consume "
    "approximately 900,000 m³ per year—an unsustainable demand in the context of Kazakhstan's water "
    "constraints. Closed-loop direct-to-chip liquid cooling reduces this consumption by three orders of "
    "magnitude to approximately 1,000 m³ (maintenance water only), while immersion cooling approaches "
    "near-zero water consumption."
)

add_figure('fig6_water_consumption.png',
           'Fig. 6. Annual water consumption comparison by cooling technology for a 35 MW facility '
           '(logarithmic scale). Traditional evaporative cooling is unsustainable given Kazakhstan\'s '
           'water projections, necessitating closed-loop or waterless alternatives.')

# Table II: Cooling Technology Comparison
add_table(
    headers=['Technology', 'WUE (L/kWh)', 'Water (m³/yr)', 'PUE Impact', 'Suitability'],
    rows=[
        ['Evaporative Tower', '1.8–2.5', '900,000', '1.10–1.20', 'Not recommended'],
        ['Hybrid (Air+Evap)', '0.5–1.0', '350,000', '1.15–1.25', 'Limited use'],
        ['Air-Cooled Chiller', '~0', '5,000', '1.20–1.35', 'Recommended'],
        ['Direct-to-Chip Liquid', '<0.1', '1,000', '1.05–1.12', 'Strongly recommended'],
        ['Immersion Cooling', '~0', '<200', '1.03–1.08', 'Optimal for AI'],
    ],
    caption='Table II. Cooling Technology Comparison for 35 MW AI Data Center Facility in Kazakhstan'
)

# ── F. Kazakhstan Energy System Context ──
add_heading_styled('F. Kazakhstan Energy System Context', level=2)

add_body(
    "Understanding Kazakhstan's evolving energy landscape is essential for long-term data center planning. "
    "As shown in Fig. 7, the current generation mix remains dominated by fossil fuels, with coal accounting "
    "for 54% and natural gas for 29% of total electricity production in 2024 [9]. Renewable energy sources "
    "(wind, solar, small hydro, biomass) collectively contribute approximately 7%, with installed renewable "
    "capacity reaching 2,904 MW across 169 projects [18]. The government's target of 15% renewable generation "
    "by 2030 requires a substantial acceleration of deployment."
)

add_figure('fig7_energy_mix.png',
           'Fig. 7. Kazakhstan electricity generation mix: current state (2024) versus the projected '
           '30% RES scenario for 2035, based on Zhakiyev et al. [10] PyPSA-KZ modeling. The transition '
           'involves significant wind and solar capacity additions while reducing coal dependency.')

add_body(
    "The grid-level carbon intensity of approximately 600 gCO₂eq/kWh [9] presents a reputational and "
    "regulatory risk for data center operators targeting global ESG compliance. However, the optimization "
    "modeling by Zhakiyev et al. [12] demonstrates a viable pathway to carbon neutrality by 2060, with "
    "intermediate milestones of −20% emissions by 2030 and −40% by 2040. For near-term deployments, "
    "power purchase agreements (PPAs) for renewable energy and carbon offset mechanisms offer pathways to "
    "mitigate the carbon footprint. The PyPSA-KZ analysis by Zhakiyev et al. [10] further identifies "
    "that a 30% RES scenario for 2035 is achievable at an annual system cost of EUR 5.08 billion, with "
    "optimal renewable placement in Turkestan and Almaty (solar) and central/eastern Kazakhstan (wind)."
)

# ── G. Regional Site Selection ──
add_heading_styled('G. Regional Site Selection', level=2)

add_body(
    "Fig. 8 presents the multi-criteria evaluation of three candidate regions. Astana achieves the "
    "highest composite score, driven by superior PUE efficiency (score: 9/10) reflecting its cold "
    "continental climate, strong network connectivity (9/10) as the national capital with established "
    "international fiber routes, and balanced performance across all criteria. Pavlodar/Ekibastuz leads "
    "decisively in power cost (10/10) due to proximity to generation capacity, but scores lower on "
    "connectivity (5/10) and renewable potential (5/10) given its coal-dominated generation base. "
    "Almaty offers the strongest renewable potential (8/10) and good connectivity (8/10) as the largest "
    "city, but its warmer climate limits PUE efficiency (5/10)."
)

add_figure('fig8_site_selection.png',
           'Fig. 8. Multi-criteria regional site selection matrix for AI data center deployment in '
           'Kazakhstan. Astana demonstrates the best overall balance, while Pavlodar/Ekibastuz leads '
           'in power cost competitiveness.')

add_body(
    "Based on this analysis, we recommend a tiered deployment strategy: Astana as the primary location "
    "for AI training hubs requiring high connectivity and balanced performance, and Pavlodar/Ekibastuz "
    "as a secondary location optimized for cost-sensitive batch processing and inference workloads. "
    "Almaty is suitable for edge computing and content delivery network (CDN) deployments leveraging "
    "its population density and regional connectivity."
)

# ════════════════════════════════════════════════════════════
# V. DISCUSSION
# ════════════════════════════════════════════════════════════
add_heading_styled('V. Discussion')

add_body(
    "Our analysis reveals that Kazakhstan possesses a structurally competitive position in the global "
    "AI infrastructure market, but one that is critically dependent on technology choices and policy "
    "decisions. The convergence of three factors—climate-enabled PUE of 1.12, electricity costs as low "
    "as $0.031/kWh with seven-year regulatory certainty, and proximity to growing Asian AI demand—creates "
    "a compelling value proposition. However, the water scarcity constraint is non-negotiable and must be "
    "addressed through mandatory technology requirements."
)

add_body(
    "The grid carbon intensity of 600 gCO₂eq/kWh represents a significant challenge for operators "
    "subject to ESG reporting requirements. A 35 MW facility operating at PUE 1.15 on the current grid "
    "would generate approximately 185,000 tonnes of CO₂ annually—equivalent to the emissions of roughly "
    "40,000 passenger vehicles. The decarbonization trajectories modeled by Zhakiyev et al. [12] suggest "
    "that grid carbon intensity will decline meaningfully only after 2030, implying that near-term "
    "deployments must rely on renewable energy PPAs and carbon credits to achieve acceptable carbon profiles."
)

add_body(
    "The cryptocurrency mining experience in Kazakhstan serves as both a cautionary tale and a relevant "
    "precedent. As documented by Zhakiyev et al. [10], bitcoin mining consumed over 1,000 MW of grid "
    "capacity (excluding unregistered operations), contributing to grid instability and prompting "
    "regulatory intervention. AI data centers must avoid repeating this pattern by ensuring dedicated "
    "power infrastructure and transparent grid impact assessments."
)

add_body(
    "The Ekibastuz Data Center Valley initiative, with its 1 GW target and dedicated power allocation "
    "from GRES stations [6], represents a structurally sound approach that aligns with our findings. "
    "The combination of station-bus pricing and proximity to coal generation (with planned transition "
    "to renewables) addresses both the economic and infrastructure requirements identified in our analysis."
)

add_body(
    "A limitation of our study is the reliance on publicly available meteorological data, which may "
    "not fully capture microclimate effects at specific facility sites. Additionally, our economic "
    "model assumes the tariff freeze will be maintained through 2032 as legislated; political or "
    "economic changes could affect this assumption. Future work should incorporate site-specific "
    "meteorological monitoring and develop stochastic models for tariff risk."
)

# ════════════════════════════════════════════════════════════
# VI. STRATEGIC RECOMMENDATIONS
# ════════════════════════════════════════════════════════════
add_heading_styled('VI. Strategic Recommendations')

add_body(
    "Based on our analysis, we propose the following evidence-based recommendations for policymakers "
    "and industry stakeholders:"
)

recommendations = [
    ("Legislative Mandate for Waterless Cooling: ",
     "Enact a regulatory requirement prohibiting the use of open-loop evaporative cooling towers for "
     "data center facilities exceeding 5 MW capacity. This is essential given the projected 50% water "
     "shortfall by 2040 and should mandate either air-cooled or closed-loop liquid cooling technologies."),

    ("Renewable Energy Integration Requirements: ",
     "Require all new data center facilities exceeding 10 MW to demonstrate a credible pathway to 50% "
     "renewable energy sourcing within five years of commissioning. The 2026–2032 tariff freeze provides "
     "a window to lock in low fossil-fuel OPEX while simultaneously investing in on-site or contracted "
     "wind and solar generation."),

    ("Waste Heat Recovery Mandate: ",
     "In Astana and other northern cities, mandate the capture and integration of data center waste heat "
     "into municipal district heating networks. A 35 MW facility at PUE 1.15 produces approximately 5 MW "
     "of recoverable thermal energy, sufficient to heat 3,000–5,000 residential units during winter months."),

    ("Grid Infrastructure Investment: ",
     "Prioritize the completion of the Atyrau–Aktobe transmission interconnection to integrate the "
     "western grid zone, enabling geographically diversified data center deployment and access to the "
     "Caspian region's market. As demonstrated by Zhakiyev et al. [10], this transmission expansion "
     "achieves 30% RES integration at marginally lower cost than scenarios without it."),
]

for label, text in recommendations:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.63)
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Times New Roman'
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.name = 'Times New Roman'

# ════════════════════════════════════════════════════════════
# VII. CONCLUSION
# ════════════════════════════════════════════════════════════
add_heading_styled('VII. Conclusion')

add_body(
    "This study presents a comprehensive techno-economic analysis of AI data center infrastructure "
    "deployment in the Republic of Kazakhstan. Our multi-factor assessment demonstrates that Kazakhstan "
    "possesses a structurally competitive position characterized by: (1) an annualized PUE of 1.12 "
    "achievable through Astana's continental climate, approaching hyperscaler-class efficiency; "
    "(2) electricity OPEX reductions of 63% versus U.S. benchmarks through co-location at generation "
    "stations, with seven-year tariff certainty; and (3) total cost of ownership potential 40–50% below "
    "Western markets."
)

add_body(
    "However, sustainable deployment is contingent upon addressing the critical water scarcity "
    "constraint through mandatory adoption of closed-loop cooling technologies, and mitigating the "
    "high grid carbon intensity through renewable energy procurement mechanisms. The regional site "
    "selection analysis identifies Astana as the optimal primary location, balancing efficiency, "
    "connectivity, and environmental factors, with Pavlodar/Ekibastuz as a complementary cost-optimized "
    "deployment zone."
)

add_body(
    "As the global demand for AI computing capacity continues its exponential growth, Kazakhstan's "
    "combination of favorable climate, competitive energy costs, and strategic geographic position—"
    "bridging European and Asian markets—positions the nation as a potentially significant participant "
    "in the international AI infrastructure landscape. The realization of this potential will depend on "
    "the implementation of the policy recommendations outlined in this study and continued investment in "
    "grid modernization and renewable energy capacity as modeled by Zhakiyev et al. [10], [12]."
)

# ════════════════════════════════════════════════════════════
# AUTHOR CONTRIBUTIONS
# ════════════════════════════════════════════════════════════
add_heading_styled('Author Contributions')
add_body_no_indent(
    "Sekenov G.: Conceptualization, Methodology, Data Curation, Formal Analysis, Visualization, "
    "Writing—Original Draft, Writing—Review and Editing."
)

# ════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ════════════════════════════════════════════════════════════
add_heading_styled('Acknowledgements')
add_body_no_indent(
    "The author acknowledges the Astana IT University for institutional support and the open-source "
    "contributions of the PyPSA-KZ modeling team led by N. Zhakiyev, whose energy system analyses "
    "provided essential context for this work."
)

# ════════════════════════════════════════════════════════════
# CONFLICT OF INTEREST
# ════════════════════════════════════════════════════════════
add_heading_styled('Conflict of Interest')
add_body_no_indent("The author declares no conflict of interest.")

# ════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════
add_heading_styled('References')

references = [
    '[1] K. M. U. Ahmed, M. H. Bollen, and M. Alvarez, "A review of data centers energy consumption and reliability modeling," IEEE Access, vol. 9, pp. 152536–152563, 2021.',
    '[2] M. Koot and F. Wijnhoven, "Usage impact on data center electricity needs: A system dynamic forecasting model," Applied Energy, vol. 291, p. 116798, 2021.',
    '[3] International Energy Agency, "Energy and AI," IEA, Paris, 2025. [Online]. Available: https://www.iea.org/reports/energy-and-ai',
    '[4] NVIDIA Corporation, "Rubin Platform," NVIDIA Newsroom, Jan. 2026. [Online]. Available: https://nvidianews.nvidia.com/news/rubin-platform-ai-supercomputer',
    '[5] Prime Minister of the Republic of Kazakhstan, "Year of Digitalization and Artificial Intelligence: The Government Has Identified Key Priorities," Jan. 2026. [Online]. Available: https://primeminister.kz',
    '[6] "Kazakhstan Advances 1 GW Data Center Valley Project in Ekibastuz," The Astana Times, Feb. 2026.',
    '[7] "Kazakhstan Proposes Seven-Year Freeze on Electricity Rate Caps," Orda.kz, Dec. 2025.',
    '[8] "Kazakhstan Expands Renewable Energy Capacity," The Astana Times, Jan. 2026.',
    '[9] Low Carbon Power, "Kazakhstan—Electricity Generation and Carbon Intensity," 2025. [Online]. Available: https://lowcarbonpower.org/region/Kazakhstan',
    '[10] N. Zhakiyev, Y. Akhmetov, R. Omirgaliyev, B. Mukatov, N. Baisakalova, S. Zhakiyeva, and B. Kazbekov, "Comprehensive scenario analyses for coal exit and renewable energy development planning of Kazakhstan using PyPSA-KZ," Engineered Science, vol. 29, p. 1085, 2024.',
    '[11] Earth.Org, "Kazakhstan\'s Water Crisis, Explained with Data," Jul. 2025. [Online]. Available: https://earth.org',
    '[12] N. Zhakiyev, A. Khamzina, S. Zhakiyeva, R. De Miglio, A. Bakdolotov, and C. Cosmi, "Optimization modelling of the decarbonization scenario of the total energy system of Kazakhstan until 2060," Energies, vol. 16, no. 13, p. 5142, 2023.',
    '[13] D. Mytton, "Data centre water consumption," npj Clean Water, vol. 4, no. 11, 2021.',
    '[14] Arctida, "Free Cooling in Arctic Data Centers," 2025. [Online]. Available: https://arctida.io/en/projects/free-cooling',
    '[15] ASHRAE Technical Committee 9.9, Thermal Guidelines for Data Processing Environments, 5th ed. Atlanta, GA: ASHRAE, 2021.',
    '[16] World Bank, "Kazakhstan Water Security Assessment," Washington, DC, 2024. [Online]. Available: https://documents.worldbank.org',
    '[17] Uptime Institute, "Global Data Center Survey Results 2024," Uptime Institute, 2024.',
    '[18] QazaqGreen, "Kazakhstan Renewable Energy: 455 MW Growth in 2025," 2025. [Online]. Available: https://www.pvknowhow.com/news/kazakhstan-renewable-energy-stunning-455-mw-growth-in-2025/',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    # Bold the reference number
    bracket_end = ref.index(']') + 1
    run_num = p.add_run(ref[:bracket_end] + ' ')
    run_num.font.size = Pt(9)
    run_num.font.name = 'Times New Roman'
    run_text = p.add_run(ref[bracket_end:].strip())
    run_text.font.size = Pt(9)
    run_text.font.name = 'Times New Roman'

# ── Save ──
doc.save(OUT)
print(f'\nPaper saved to: {OUT}')
print(f'Total figures: 8')
print(f'Total tables: 2')
print(f'Total references: {len(references)}')
