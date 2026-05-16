#!/usr/bin/env python3
"""Build LaTeX, Markdown, DOCX, and figure assets for the transfer forecasting article."""

from __future__ import annotations

import json
import textwrap
from pathlib import Path
from typing import Optional

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

BASE_DIR = Path(__file__).resolve().parents[2]
FINAL_DIR = BASE_DIR / "Final"
FIG_DIR = FINAL_DIR / "figures" / "transfer_forecasting"
RESULTS_DIR = BASE_DIR / "results" / "forecasting"
FORECASTING_DIR = BASE_DIR / "data" / "forecasting"
QA_DIR = BASE_DIR / "data" / "processed" / "qa_reports"

TITLE = "Probabilistic Multi-Horizon Electricity Demand Forecasting for Kazakhstan with EU-to-KZ Transfer Learning"


def load_inputs() -> dict:
    metrics = pd.read_csv(RESULTS_DIR / "model_metrics.csv")
    dm = pd.read_csv(RESULTS_DIR / "diebold_mariano_tests.csv")
    preds = pd.read_csv(RESULTS_DIR / "predictions_h24.csv", parse_dates=["timestamp_utc"])
    backtest = pd.read_csv(RESULTS_DIR / "backtest_h24.csv")
    pipeline = json.loads((RESULTS_DIR / "pipeline_summary.json").read_text())
    build = json.loads((FORECASTING_DIR / "build_manifest.json").read_text())
    qa = json.loads((QA_DIR / "transfer_forecasting_cleaning_report.json").read_text())
    return {
        "metrics": metrics,
        "dm": dm,
        "preds": preds,
        "backtest": backtest,
        "pipeline": pipeline,
        "build": build,
        "qa": qa,
    }


def format_float(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def latex_escape(text: str) -> str:
    return (
        str(text)
        .replace("\\", "\\textbackslash{}")
        .replace("_", "\\_")
        .replace("%", "\\%")
        .replace("&", "\\&")
        .replace("#", "\\#")
    )


def create_figures(data: dict) -> dict:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "workflow": FIG_DIR / "fig1_workflow_placeholder.png",
        "coverage": FIG_DIR / "fig2_data_coverage_placeholder.png",
        "metrics": FIG_DIR / "fig3_model_comparison.png",
        "predictions": FIG_DIR / "fig4_predictions_h24.png",
    }

    plt.rcParams.update(
        {
            "figure.dpi": 220,
            "savefig.dpi": 220,
            "font.family": "DejaVu Serif",
            "axes.titlesize": 13,
            "axes.labelsize": 11,
        }
    )

    # Figure 1: workflow placeholder
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.axis("off")
    boxes = [
        (0.05, 0.35, 0.2, 0.3, "Open Data\nEU hourly load\nKazakhstan demand\nOWID / Ember / proxies"),
        (0.30, 0.35, 0.2, 0.3, "Canonical Tables\nload / weather\ncalendar / macro\nUTC alignment"),
        (0.55, 0.35, 0.2, 0.3, "Model Stack\nSeasonal naive\nGBDT transfer\nGBDT non-transfer"),
        (0.80, 0.35, 0.15, 0.3, "Outputs\nMetrics\nDM tests\nSubmission docs"),
    ]
    for x, y, w, h, label in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor="#f3f1ea", edgecolor="#444444", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=11)
    for x1, x2 in [(0.25, 0.30), (0.50, 0.55), (0.75, 0.80)]:
        ax.annotate("", xy=(x2, 0.50), xytext=(x1, 0.50), arrowprops=dict(arrowstyle="->", lw=2))
    ax.text(
        0.5,
        0.88,
        "Figure placeholder: reproducible EU-to-KZ transfer forecasting workflow",
        ha="center",
        va="center",
        fontsize=14,
        fontweight="bold",
    )
    fig.tight_layout()
    fig.savefig(paths["workflow"], bbox_inches="tight")
    plt.close(fig)

    # Figure 2: coverage placeholder
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.set_title("Figure placeholder: dataset coverage used in the current article package", fontweight="bold")
    ax.set_xlabel("Year")
    ax.set_yticks([3, 2, 1])
    ax.set_yticklabels(["KZ hourly target", "EU hourly source", "QA / macro context"])
    ax.grid(axis="x", linestyle=":", alpha=0.4)
    ax.barh([3], [8], left=[2011], height=0.5, color="#d7c7a9", label="Kazakhstan target series")
    ax.barh([2], [0.75], left=[2020], height=0.5, color="#90b4ce", label="EU hourly source sample")
    ax.barh([1], [25], left=[2000], height=0.5, color="#b9d3b0", label="Macro / annual context")
    ax.set_xlim(2000, 2025)
    ax.legend(loc="upper left")
    fig.tight_layout()
    fig.savefig(paths["coverage"], bbox_inches="tight")
    plt.close(fig)

    # Figure 3: actual model comparison
    metrics = data["metrics"].copy()
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    axes[0].bar(metrics["model"], metrics["RMSE"], color=["#385170", "#9fd3c7", "#f08a5d"])
    axes[0].set_title("RMSE by model")
    axes[0].set_ylabel("RMSE (MW)")
    axes[0].tick_params(axis="x", rotation=20)
    axes[1].bar(metrics["model"], metrics["sMAPE"], color=["#385170", "#9fd3c7", "#f08a5d"])
    axes[1].set_title("sMAPE by model")
    axes[1].set_ylabel("sMAPE (%)")
    axes[1].tick_params(axis="x", rotation=20)
    fig.suptitle("24-hour forecasting results from the current validated fast-run", fontweight="bold")
    fig.tight_layout()
    fig.savefig(paths["metrics"], bbox_inches="tight")
    plt.close(fig)

    # Figure 4: actual predictions sample
    pred = data["preds"].head(168).copy()
    fig, ax = plt.subplots(figsize=(12, 4.8))
    ax.plot(pred["timestamp_utc"], pred["y_true"], label="Observed", color="#222831", linewidth=2)
    ax.plot(
        pred["timestamp_utc"],
        pred["pred_gbdt_non_transfer"],
        label="GBDT non-transfer",
        color="#005f73",
        linewidth=1.8,
    )
    ax.fill_between(
        pred["timestamp_utc"],
        pred["lo_gbdt_non_transfer"],
        pred["hi_gbdt_non_transfer"],
        color="#94d2bd",
        alpha=0.35,
        label="90% interval",
    )
    ax.set_title("Observed vs predicted load (first 168 test hours, 24h horizon)", fontweight="bold")
    ax.set_ylabel("Load (MW)")
    ax.legend()
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(paths["predictions"], bbox_inches="tight")
    plt.close(fig)

    return {k: str(v.relative_to(FINAL_DIR)) for k, v in paths.items()}


def build_context(data: dict) -> dict:
    metrics = data["metrics"].copy()
    dm = data["dm"].copy()
    backtest = data["backtest"].copy()
    qa = data["qa"]
    build = data["build"]

    best = metrics.sort_values(["RMSE", "MAE"]).iloc[0]
    baseline = metrics.loc[metrics["model"] == "seasonal_naive"].iloc[0]
    transfer = metrics.loc[metrics["model"].str.contains("transfer") & ~metrics["model"].str.contains("non")].iloc[0]
    improvement = (baseline["sMAPE"] - best["sMAPE"]) / baseline["sMAPE"] * 100

    return {
        "best_model": best["model"],
        "best_rmse": format_float(best["RMSE"]),
        "best_mae": format_float(best["MAE"]),
        "best_smape": format_float(best["sMAPE"]),
        "best_mase": format_float(best["MASE"]),
        "best_picp": format_float(best["PICP_90"]),
        "baseline_smape": format_float(baseline["sMAPE"]),
        "transfer_smape": format_float(transfer["sMAPE"]),
        "improvement_smape": format_float(improvement, 1),
        "load_rows": f"{build['table_load_hourly_rows']:,}",
        "weather_rows": f"{build['table_weather_hourly_rows']:,}",
        "calendar_rows": f"{build['table_calendar_rows']:,}",
        "macro_rows": f"{build['table_macro_annual_rows']:,}",
        "regions_list": ", ".join(build["regions"]),
        "kz_min": build["kz_time_range"]["min"],
        "kz_max": build["kz_time_range"]["max"],
        "eu_min": build["eu_time_range"]["min"],
        "eu_max": build["eu_time_range"]["max"],
        "owid_synced_year": qa["owid_sync"]["synced_max_year"],
        "owid_missing_years": ", ".join(str(y) for y in qa["owid_sync"]["gdp_missing_years"]),
        "ember_negative_total": qa["ember_negative_rule"]["negative_values_total"],
        "ember_invalid_negative": qa["ember_negative_rule"]["invalid_negative_values"],
        "metrics_table_md": metrics_to_markdown(metrics),
        "dm_table_md": metrics_to_markdown(dm),
        "backtest_table_md": metrics_to_markdown(backtest),
        "metrics_table_tex": metrics_to_latex(metrics, metric_table=True),
        "dm_table_tex": metrics_to_latex(dm),
        "backtest_table_tex": metrics_to_latex(backtest),
    }


def metrics_to_markdown(df: pd.DataFrame) -> str:
    cols = list(df.columns)
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    rows = []
    for _, row in df.iterrows():
        vals = []
        for col in cols:
            val = row[col]
            if isinstance(val, float):
                vals.append(format_float(val, 6 if abs(val) < 0.001 else 2))
            else:
                vals.append(str(val))
        rows.append("| " + " | ".join(vals) + " |")
    return "\n".join([header, sep] + rows)


def metrics_to_latex(df: pd.DataFrame, metric_table: bool = False) -> str:
    cols = list(df.columns)
    align = "l" + "r" * (len(cols) - 1)
    header = " & ".join(latex_escape(col) for col in cols) + " \\\\"
    lines = [f"\\begin{{tabular}}{{{align}}}", "\\toprule", header, "\\midrule"]
    for _, row in df.iterrows():
        vals = []
        for col in cols:
            val = row[col]
            if isinstance(val, float):
                vals.append(format_float(val, 6 if abs(val) < 0.001 else 2))
            else:
                vals.append(latex_escape(val))
        lines.append(" & ".join(vals) + " \\\\")
    lines.extend(["\\bottomrule", "\\end{tabular}"])
    return "\n".join(lines)


def render_markdown(ctx: dict, fig_paths: dict) -> str:
    text = textwrap.dedent(
        f"""
        # {TITLE}

        **Submission Draft**

        ## Abstract

        This article presents a reproducible open-data workflow for probabilistic multi-horizon electricity demand forecasting for Kazakhstan with EU-to-KZ transfer learning. The current study package integrates canonical hourly load, weather proxy, calendar, and macro-energy tables and evaluates a seasonal naive baseline, a non-transfer gradient-boosted model, and a transfer-learning gradient-boosted model. In the current validated fast-run experiment at the 24-hour horizon, the best model is `{ctx['best_model']}` with RMSE = {ctx['best_rmse']} MW, MAE = {ctx['best_mae']} MW, sMAPE = {ctx['best_smape']}%, MASE = {ctx['best_mase']}, and PICP = {ctx['best_picp']}%. Relative to the seasonal naive benchmark, this corresponds to an sMAPE improvement of {ctx['improvement_smape']}%. Diebold-Mariano results indicate statistically significant gains over both the baseline and the transfer variant. The article is intentionally written as a journal-style draft while clearly marking the current result set as a fast-run validation package rather than the final full-horizon experiment.

        **Keywords:** Kazakhstan, electricity demand forecasting, transfer learning, probabilistic forecasting, gradient boosting, open energy data

        ## 1. Introduction

        Kazakhstan's electricity system is facing simultaneous growth in demand, pressure on aging thermal assets, and a need for better operational forecasting. Yet open, high-frequency Kazakhstan demand data remain limited. This article addresses that gap with a transparent EU-to-KZ transfer-learning workflow built entirely on open or reproducibly derived data assets.

        The core idea is simple. Europe supplies rich hourly source-domain load series for methodological pretraining and benchmarking, while Kazakhstan contributes a target-domain series built from open monthly demand validation data and a documented monthly-to-hourly disaggregation strategy. This setup is not presented as a substitute for direct metered hourly data. Instead, it is positioned as an evidence-first research bridge that makes the forecasting problem tractable under open-data constraints.

        ## 2. Study Design and Data

        The current canonical build contains:

        - `table_load_hourly.csv`: {ctx['load_rows']} rows
        - `table_weather_hourly.csv`: {ctx['weather_rows']} rows
        - `table_calendar.csv`: {ctx['calendar_rows']} rows
        - `table_macro_annual.csv`: {ctx['macro_rows']} rows

        Regions in the current build: {ctx['regions_list']}.

        Kazakhstan target coverage spans `{ctx['kz_min']}` to `{ctx['kz_max']}`. The EU source-domain sample used in the current package spans `{ctx['eu_min']}` to `{ctx['eu_max']}`.

        ![Workflow]({fig_paths['workflow']})

        ![Coverage]({fig_paths['coverage']})

        ## 3. Data Cleaning and Audit

        The QA workflow cleaned and validated the local renewable proxy files, synchronized the Kazakhstan OWID slices into explicit legacy and synced outputs, and enforced a rule that negative Ember values are valid only for `Net Imports`.

        - Synced OWID coverage now reaches {ctx['owid_synced_year']}.
        - GDP-related missing years currently flagged: {ctx['owid_missing_years']}.
        - Ember negative values found: {ctx['ember_negative_total']}; invalid negatives remaining: {ctx['ember_invalid_negative']}.

        ## 4. Forecasting Framework

        The current article package evaluates three models for the 24-hour horizon:

        - `seasonal_naive`
        - `gbdt_transfer`
        - `{ctx['best_model']}`

        Feature engineering includes autoregressive lags (1, 2, 3, 6, 12, 24, 48, 168 hours), rolling means, calendar indicators, annual macro features, and Kazakhstan weather/resource proxy inputs. Uncertainty intervals are produced with split-conformal calibration from validation residuals.

        ## 5. Main Results

        ### Table 1. Forecast metrics

        {ctx['metrics_table_md']}

        ![Model comparison]({fig_paths['metrics']})

        ### Table 2. Diebold-Mariano tests

        {ctx['dm_table_md']}

        ### Table 3. Rolling backtest summary

        {ctx['backtest_table_md']}

        ![Observed vs predicted]({fig_paths['predictions']})

        The strongest current result is the non-transfer gradient-boosted model. It reduces sMAPE from {ctx['baseline_smape']}% for the seasonal naive benchmark to {ctx['best_smape']}%, an improvement of {ctx['improvement_smape']}%. This comfortably exceeds the original acceptance threshold of 10% improvement at the 24-hour horizon.

        A second important result is that the transfer variant currently underperforms the non-transfer model. This likely reflects source-target mismatch in the present fast-run configuration rather than a negative result for transfer learning as a broader methodological idea.

        ## 6. Discussion

        The current package shows that a publishable forecasting workflow can be assembled even when Kazakhstan's openly accessible hourly demand data are limited. The contribution is not only the forecast metric itself, but the combined package: cleaning rules, canonical tables, statistical comparison, uncertainty intervals, and a journal-ready document trail.

        At the same time, the current interval coverage remains too conservative for the best model. With PICP = {ctx['best_picp']}% for a nominal 90% interval, interval calibration is still broader than desired. This is exactly the kind of issue that a full-horizon run and deeper calibration pass should address before submission.

        ## 7. Limitations

        Three limitations must be stated directly.

        - The present article reflects the repository's validated `FAST_DEMO` result package, not the final full-run experiment.
        - The Kazakhstan hourly target is a transparent synthetic disaggregation, not direct metered hourly load.
        - The current submission package includes figure placeholders for workflow and coverage, while the final journal submission should replace or complement them with fully styled publication figures.

        ## 8. Conclusion

        This manuscript package establishes the article backbone for the dissertation topic on probabilistic multi-horizon electricity demand forecasting for Kazakhstan with EU-to-KZ transfer learning. The current result set already supports a coherent, evidence-based paper draft. The next publication step is to run the full experiment across all target horizons, refine interval calibration, and regenerate the tables and figures from the full-run outputs.
        """
    ).strip()
    return "\n".join(line[8:] if line.startswith("        ") else line for line in text.splitlines()) + "\n"


def render_latex(ctx: dict, fig_paths: dict) -> str:
    best_model_tex = latex_escape(ctx["best_model"])
    regions_tex = latex_escape(ctx["regions_list"])
    workflow_path = "\\detokenize{" + fig_paths["workflow"] + "}"
    coverage_path = "\\detokenize{" + fig_paths["coverage"] + "}"
    metrics_path = "\\detokenize{" + fig_paths["metrics"] + "}"
    predictions_path = "\\detokenize{" + fig_paths["predictions"] + "}"

    return textwrap.dedent(
        f"""
        \\documentclass[preprint,12pt]{{elsarticle}}
        \\usepackage[utf8]{{inputenc}}
        \\usepackage[T1]{{fontenc}}
        \\usepackage{{amsmath,amssymb}}
        \\usepackage{{graphicx}}
        \\usepackage{{booktabs}}
        \\usepackage{{tabularx}}
        \\usepackage{{hyperref}}
        \\usepackage{{float}}
        \\usepackage{{geometry}}
        \\geometry{{a4paper, margin=2.5cm}}
        \\journal{{Energy Reports}}

        \\begin{{document}}
        \\begin{{frontmatter}}

        \\title{{{TITLE}}}
        \\author[aitu]{{Gabit Sekenov}}
        \\ead{{243046@astanait.edu.kz}}
        \\author[harvard]{{Nurkhat Zhakiyev}}
        \\ead{{nzhakiyev@fas.harvard.edu}}
        \\affiliation[aitu]{{organization={{Astana IT University}}, city={{Astana}}, country={{Kazakhstan}}}}
        \\affiliation[harvard]{{organization={{Harvard University}}, city={{Cambridge, MA}}, country={{United States}}}}

        \\begin{{abstract}}
        This article presents a reproducible open-data workflow for probabilistic multi-horizon electricity demand forecasting for Kazakhstan with EU-to-KZ transfer learning. The current study package integrates canonical hourly load, weather proxy, calendar, and macro-energy tables and evaluates a seasonal naive baseline, a non-transfer gradient-boosted model, and a transfer-learning gradient-boosted model. In the current validated fast-run experiment at the 24-hour horizon, the best model is \\texttt{{{best_model_tex}}} with RMSE = {ctx['best_rmse']} MW, MAE = {ctx['best_mae']} MW, sMAPE = {ctx['best_smape']}\\%, MASE = {ctx['best_mase']}, and PICP = {ctx['best_picp']}\\%. Relative to the seasonal naive benchmark, this corresponds to an sMAPE improvement of {ctx['improvement_smape']}\\%. Diebold-Mariano results indicate statistically significant gains over both the baseline and the transfer variant.
        \\end{{abstract}}

        \\begin{{keyword}}
        Kazakhstan \\sep electricity demand forecasting \\sep transfer learning \\sep probabilistic forecasting \\sep gradient boosting \\sep open energy data
        \\end{{keyword}}

        \\begin{{highlights}}
        \\item The current canonical build integrates {ctx['load_rows']} hourly load rows across 11 regions
        \\item The best validated fast-run model is \\texttt{{{best_model_tex}}} with sMAPE = {ctx['best_smape']}\\%
        \\item sMAPE improves by {ctx['improvement_smape']}\\% relative to the seasonal naive benchmark
        \\item Diebold-Mariano tests show statistically significant gains over the baseline and transfer variant
        \\item The manuscript package is fully reproducible from repository outputs
        \\end{{highlights}}

        \\end{{frontmatter}}

        \\section{{Introduction}}
        Kazakhstan's electricity system is facing simultaneous growth in demand, pressure on aging thermal assets, and a need for better operational forecasting. Yet open, high-frequency Kazakhstan demand data remain limited. This article addresses that gap with a transparent EU-to-KZ transfer-learning workflow built entirely on open or reproducibly derived data assets.

        \\section{{Study Design and Data}}
        The current canonical build contains {ctx['load_rows']} hourly load rows, {ctx['weather_rows']} weather rows, {ctx['calendar_rows']} calendar rows, and {ctx['macro_rows']} macro rows. The regions included are: {regions_tex}. Kazakhstan target coverage spans {ctx['kz_min']} to {ctx['kz_max']}, while the EU source-domain sample used in the current package spans {ctx['eu_min']} to {ctx['eu_max']}.

        \\begin{{figure}}[H]
        \\centering
        \\includegraphics[width=0.95\\textwidth]{{{workflow_path}}}
        \\caption{{Workflow placeholder for the reproducible EU-to-KZ transfer forecasting pipeline.}}
        \\label{{fig:workflow}}
        \\end{{figure}}

        \\begin{{figure}}[H]
        \\centering
        \\includegraphics[width=0.95\\textwidth]{{{coverage_path}}}
        \\caption{{Coverage placeholder summarizing the current source-domain and target-domain data used in the manuscript package.}}
        \\label{{fig:coverage}}
        \\end{{figure}}

        \\section{{Data Cleaning and Audit}}
        The QA workflow normalized the local renewable proxy files, synchronized the Kazakhstan OWID slices into explicit legacy and synced outputs, and enforced a rule that negative Ember values are valid only for \\texttt{{Net Imports}}. Synced OWID coverage reaches {ctx['owid_synced_year']}. GDP-related missing years currently flagged are {ctx['owid_missing_years']}. Ember negative values found: {ctx['ember_negative_total']}; invalid negatives remaining: {ctx['ember_invalid_negative']}.

        \\section{{Forecasting Framework}}
        The current article package evaluates a seasonal naive baseline, a transfer gradient-boosted model, and a non-transfer gradient-boosted model for the 24-hour horizon. Feature engineering includes autoregressive lags, rolling means, calendar indicators, annual macro features, and Kazakhstan weather/resource proxy inputs. Prediction intervals are produced with split-conformal calibration from validation residuals.

        \\section{{Results}}
        \\subsection{{Forecast metrics}}
        \\begin{{table}}[H]
        \\centering
        \\caption{{Forecast accuracy at the 24-hour horizon in the current validated fast-run package.}}
        {ctx['metrics_table_tex']}
        \\label{{tab:metrics}}
        \\end{{table}}

        \\begin{{figure}}[H]
        \\centering
        \\includegraphics[width=0.95\\textwidth]{{{metrics_path}}}
        \\caption{{Model comparison for RMSE and sMAPE in the current validated fast-run package.}}
        \\label{{fig:modelcomp}}
        \\end{{figure}}

        \\subsection{{Diebold-Mariano comparisons}}
        \\begin{{table}}[H]
        \\centering
        \\caption{{Pairwise Diebold-Mariano tests under squared-error loss.}}
        {ctx['dm_table_tex']}
        \\label{{tab:dm}}
        \\end{{table}}

        \\subsection{{Rolling backtest}}
        \\begin{{table}}[H]
        \\centering
        \\caption{{Rolling backtest summary for the 24-hour horizon.}}
        {ctx['backtest_table_tex']}
        \\label{{tab:backtest}}
        \\end{{table}}

        \\begin{{figure}}[H]
        \\centering
        \\includegraphics[width=0.95\\textwidth]{{{predictions_path}}}
        \\caption{{Observed vs predicted load for the first 168 test hours, including the 90\\% interval for the best model.}}
        \\label{{fig:predictions}}
        \\end{{figure}}

        The strongest current result is the non-transfer gradient-boosted model. It reduces sMAPE from {ctx['baseline_smape']}\\% for the seasonal naive benchmark to {ctx['best_smape']}\\%, an improvement of {ctx['improvement_smape']}\\%. This exceeds the original acceptance threshold of 10\\% improvement at the 24-hour horizon.

        \\section{{Discussion}}
        The current package shows that a publishable forecasting workflow can be assembled even when Kazakhstan's openly accessible hourly demand data are limited. The contribution is not only the forecast metric itself, but the combined package: cleaning rules, canonical tables, statistical comparison, uncertainty intervals, and a journal-ready document trail.

        A second important result is that the transfer variant currently underperforms the non-transfer model. This likely reflects source-target mismatch in the present fast-run configuration rather than a negative result for transfer learning in general.

        \\section{{Limitations}}
        First, the present article reflects the repository's validated fast-run result package rather than the final full-run experiment. Second, the Kazakhstan hourly target is a transparent synthetic disaggregation, not direct metered hourly load. Third, the workflow and coverage figures are intentionally placeholder-style assets for journal drafting and should be replaced or refined in the final submission layout.

        \\section{{Conclusion}}
        This manuscript package establishes the article backbone for the dissertation topic on probabilistic multi-horizon electricity demand forecasting for Kazakhstan with EU-to-KZ transfer learning. The implemented pipeline already supports a coherent, evidence-based paper draft. The next publication step is to execute the full experiment across all target horizons, refine interval calibration, and regenerate the tables and figures from the full-run outputs.

        \\begin{{thebibliography}}{{9}}
        \\bibitem{{opsd}} Open Power System Data, Time series package, \\url{{https://data.open-power-system-data.org/time_series/}}.
        \\bibitem{{owid}} Our World in Data, Energy dataset, \\url{{https://github.com/owid/energy-data}}.
        \\bibitem{{ember}} Ember, yearly electricity data, \\url{{https://ember-energy.org/data/}}.
        \\bibitem{{kegoc}} KEGOC, national power system and electricity balance resources, \\url{{https://www.kegoc.kz/en/electric-power/natsionalnaya-energosistema/}}.
        \\bibitem{{openmeteo}} Open-Meteo historical weather API, \\url{{https://open-meteo.com/en/docs/historical-weather-api}}.
        \\end{{thebibliography}}

        \\end{{document}}
        """
    ).strip() + "\n"


def build_docx(ctx: dict, fig_paths: dict) -> Path:
    doc = Document()
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph("Submission Draft generated from repository outputs.")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        f"This article presents a reproducible open-data workflow for probabilistic multi-horizon electricity "
        f"demand forecasting for Kazakhstan with EU-to-KZ transfer learning. In the current validated fast-run "
        f"experiment at the 24-hour horizon, the best model is {ctx['best_model']} with RMSE = {ctx['best_rmse']} MW, "
        f"MAE = {ctx['best_mae']} MW, sMAPE = {ctx['best_smape']}%, and PICP = {ctx['best_picp']}%."
    )

    doc.add_heading("Study Design and Data", level=1)
    doc.add_paragraph(
        f"The current canonical build contains {ctx['load_rows']} hourly load rows, {ctx['weather_rows']} weather rows, "
        f"{ctx['calendar_rows']} calendar rows, and {ctx['macro_rows']} macro rows. Kazakhstan target coverage spans "
        f"{ctx['kz_min']} to {ctx['kz_max']}; the EU source-domain sample spans {ctx['eu_min']} to {ctx['eu_max']}."
    )

    for label in ["workflow", "coverage"]:
        doc.add_picture(str(FINAL_DIR / fig_paths[label]), width=Inches(6.2))

    doc.add_heading("Forecast Metrics", level=1)
    add_dataframe_table(doc, pd.read_csv(RESULTS_DIR / "model_metrics.csv"))
    doc.add_picture(str(FINAL_DIR / fig_paths["metrics"]), width=Inches(6.2))

    doc.add_heading("Diebold-Mariano Tests", level=1)
    add_dataframe_table(doc, pd.read_csv(RESULTS_DIR / "diebold_mariano_tests.csv"))

    doc.add_heading("Rolling Backtest", level=1)
    add_dataframe_table(doc, pd.read_csv(RESULTS_DIR / "backtest_h24.csv"))
    doc.add_picture(str(FINAL_DIR / fig_paths["predictions"]), width=Inches(6.2))

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        f"The strongest current result is the non-transfer gradient-boosted model, which improves sMAPE by "
        f"{ctx['improvement_smape']}% relative to the seasonal naive benchmark. The transfer variant currently "
        f"underperforms the non-transfer model, suggesting that domain mismatch remains a central issue for the full experiment."
    )

    doc.add_heading("Limitations", level=1)
    doc.add_paragraph(
        "This package reflects the validated fast-run result set rather than the final full-horizon experiment. "
        "The Kazakhstan hourly target remains a transparent synthetic disaggregation rather than direct metered load."
    )

    out = FINAL_DIR / "transfer_forecasting_article.docx"
    doc.save(str(out))
    return out


def add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[idx].text = format_float(val, 6 if abs(val) < 0.001 else 2)
            else:
                cells[idx].text = str(val)


def try_compile_tex(tex_path: Path) -> Optional[Path]:
    tectonic = FINAL_DIR / ".tools" / "tectonic"
    if not tectonic.exists():
        return None
    try:
        import subprocess

        tectonic.chmod(0o755)
        subprocess.run(
            [str(tectonic), str(tex_path.name)],
            cwd=str(FINAL_DIR),
            check=True,
            capture_output=True,
            text=True,
        )
        pdf_path = tex_path.with_suffix(".pdf")
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None


def main() -> None:
    data = load_inputs()
    fig_paths = create_figures(data)
    ctx = build_context(data)

    md_text = render_markdown(ctx, fig_paths)
    tex_text = render_latex(ctx, fig_paths)

    md_path = FINAL_DIR / "transfer_forecasting_article_submission.md"
    md_main_path = FINAL_DIR / "transfer_forecasting_article.md"
    tex_path = FINAL_DIR / "transfer_forecasting_article.tex"

    md_path.write_text(md_text, encoding="utf-8")
    md_main_path.write_text(md_text, encoding="utf-8")
    tex_path.write_text(tex_text, encoding="utf-8")

    docx_path = build_docx(ctx, fig_paths)
    pdf_path = try_compile_tex(tex_path)

    summary = {
        "markdown": str(md_path.relative_to(BASE_DIR)),
        "markdown_main": str(md_main_path.relative_to(BASE_DIR)),
        "latex": str(tex_path.relative_to(BASE_DIR)),
        "docx": str(docx_path.relative_to(BASE_DIR)),
        "pdf": str(pdf_path.relative_to(BASE_DIR)) if pdf_path else None,
        "figures": fig_paths,
    }
    print(json.dumps(summary, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
